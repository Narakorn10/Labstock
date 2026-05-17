from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sop_doc():
    doc = Document()

    # Title
    title = doc.add_heading('คู่มือการใช้งานระบบควบคุมคลังน้ำยา (Lab Smart System)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Intro
    p = doc.add_paragraph()
    p.add_run('ระบบ Lab Smart System ถูกพัฒนาขึ้นเพื่อใช้ในการจัดการคลังน้ำยาในห้องปฏิบัติการ ช่วยลดความผิดพลาดในการจัดการสต๊อก และมีระบบแจ้งเตือนน้ำยาใกล้หมดผ่าน LINE Notify').italic = True

    # Section 1: การตั้งค่าระบบครั้งแรก (Setup)
    doc.add_heading('1. การเริ่มใช้งานระบบครั้งแรก (Setup)', level=1)
    doc.add_paragraph('เมื่อเปิดโปรแกรมครั้งแรก ให้ไปที่แท็บ "Master" หรือ "ข้อมูล" แล้วกดปุ่ม "Setup DB" ระบบจะสร้างแผ่นงานที่จำเป็นใน Google Sheets ดังนี้:')
    setup_list = [
        'MasterData: เก็บข้อมูลหลักของน้ำยา',
        'Inventory: เก็บยอดคงเหลือราย Lot',
        'Logs: เก็บประวัติการทำรายการทั้งหมด',
        'Settings: เก็บค่าตัวเลือกต่างๆ เช่น ประเภทน้ำยา, ชื่อเครื่องมือ'
    ]
    for item in setup_list:
        doc.add_paragraph(item, style='List Bullet')

    # Section 2: การจัดการข้อมูลหลัก (Master Data)
    doc.add_heading('2. การจัดการข้อมูลหลัก (Master Data)', level=1)
    doc.add_paragraph('เป็นการบันทึกข้อมูลน้ำยาที่ห้องแล็บมีใช้งาน เพื่อให้ระบบรู้จักน้ำยานั้นๆ')
    master_steps = [
        'ไปที่แท็บ "Master" กดปุ่มเครื่องหมายบวก (+)',
        'กรอก Item ID (รหัสอ้างอิงภายใน) และ Barcode (รหัสที่อยู่บนขวด)',
        'ระบุชื่อน้ำยา, หน่วยนับ, และจุดแจ้งเตือน (Min Alert) เมื่อยอดเหลือต่ำกว่าจุดนี้ระบบจะเตือน',
        'ระบุ Weekly Target (เป้าหมายสต๊อกหน้างาน) เพื่อใช้ในระบบคำนวณยอดเบิกอัตโนมัติ'
    ]
    for step in master_steps:
        doc.add_paragraph(step, style='List Bullet')

    # Section 3: การรับน้ำยาเข้าคลัง (Receive)
    doc.add_heading('3. การรับน้ำยาเข้าคลัง (Receive)', level=1)
    doc.add_paragraph('ใช้เมื่อมีน้ำยาใหม่ส่งมาจากพัสดุหรือบริษัท:')
    receive_steps = [
        'ไปที่แท็บ "รับเข้า" (Receive)',
        'กดปุ่ม "เปิดกล้องสแกน" หรือพิมพ์ชื่อ/รหัสในช่องค้นหา',
        'ระบุ Lot No. และวันหมดอายุ (EXP) ให้ถูกต้อง',
        'กรอกจำนวนที่รับเข้า แล้วกด "เพิ่มลงรายการ"',
        'ตรวจสอบรายการในตะกร้า แล้วกด "ยืนยันการบันทึกทั้งหมด"'
    ]
    for step in receive_steps:
        doc.add_paragraph(step, style='List Bullet')

    # Section 4: การเบิกน้ำยาไปใช้งาน (Dispense)
    doc.add_heading('4. การเบิกน้ำยาไปใช้งาน (Dispense)', level=1)
    doc.add_paragraph('ใช้เมื่อต้องการนำน้ำยาออกจากตู้เย็นหลักไปใช้ที่หน้างาน:')
    dispense_steps = [
        'ไปที่แท็บ "เบิกจ่าย" (Dispense)',
        'สแกนหรือค้นหาน้ำยาที่ต้องการ',
        'ระบบจะเลือก Lot ที่หมดอายุก่อนมาให้โดยอัตโนมัติ (หลักการ FEFO - First Expired First Out)',
        'กรอกจำนวนที่จะเบิก แล้วกด "เพิ่มลงรายการ" และ "ยืนยันการบันทึก"'
    ]
    for step in dispense_steps:
        doc.add_paragraph(step, style='List Bullet')

    # Section 5: การนับสต๊อกหน้างาน (Count & Auto-Dispense)
    doc.add_heading('5. การนับสต๊อกหน้างาน (Count)', level=1)
    doc.add_paragraph('ฟีเจอร์พิเศษสำหรับช่วยคำนวณยอดเบิกให้พอดีกับเป้าหมาย:')
    count_steps = [
        'ไปที่แท็บ "นับหน้างาน" (Count)',
        'กรอกจำนวนน้ำยาที่เหลืออยู่ที่หน้าตู้เย็น/หน้าเครื่อง',
        'ระบบจะเปรียบเทียบกับเป้าหมาย (Weekly Target) ที่ตั้งไว้ใน Master Data',
        'หากยอดเหลือน้อยกว่าเป้าหมาย จะมีปุ่ม "นำยอดไปเบิก" ปรากฏขึ้น',
        'กดปุ่มเพื่อส่งยอดที่ต้องเบิกเพิ่มไปยังแท็บเบิกจ่ายโดยไม่ต้องคำนวณเอง'
    ]
    for step in count_steps:
        doc.add_paragraph(step, style='List Bullet')

    # Section 6: รายงานและประวัติ (Reports & Logs)
    doc.add_heading('6. รายงานและประวัติ (Reports & Logs)', level=1)
    doc.add_paragraph('ตรวจสอบความเคลื่อนไหวและยอดคงเหลือ:')
    report_list = [
        'Dashboard: ดูภาพรวมยอดคงเหลือ ถ้าน้ำยาไหนต่ำกว่าจุดเตือนจะเป็นตัวหนังสือสีแดง',
        'Report: สามารถเลือกออกรายงานเป็น Excel, PDF หรือคัดลอกส่งเข้ากลุ่ม LINE ได้ที่ปุ่ม "ออกรายงาน"',
        'Logs: ดูประวัติย้อนหลังว่าใคร รับ/เบิก อะไรไป เมื่อไหร่'
    ]
    for item in report_list:
        doc.add_paragraph(item, style='List Bullet')

    # Footer
    doc.add_paragraph('\n--- จัดทำเพื่อพัฒนาการทำงานในห้องปฏิบัติการ ---').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save
    filename = 'SOP_Lab_Smart_System.docx'
    doc.save(filename)
    print(f"SOP created: {filename}")

if __name__ == "__main__":
    create_sop_doc()
